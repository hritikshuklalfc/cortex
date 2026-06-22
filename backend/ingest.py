import os
import json
from langchain_community.document_loaders import PyMuPDFLoader, CSVLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LCDocument
from dotenv import load_dotenv

load_dotenv()

CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "./chroma_db")

# Initialize embeddings
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# Initialize ChromaDB
vector_store = Chroma(
    collection_name="industrial_knowledge",
    embedding_function=embeddings,
    persist_directory=CHROMA_PERSIST_DIR
)


def load_docx(file_path: str) -> list:
    """Load a DOCX file and return LangChain Documents."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for DOCX files. Install with: pip install python-docx")
    
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    
    content = "\n".join(full_text)
    return [LCDocument(page_content=content, metadata={"source": file_path})]


def load_json(file_path: str) -> list:
    """Load a JSON file and return LangChain Documents."""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # Convert JSON to readable text
    if isinstance(data, list):
        content = "\n\n".join([json.dumps(item, indent=2) for item in data])
    else:
        content = json.dumps(data, indent=2)
    
    return [LCDocument(page_content=content, metadata={"source": file_path})]


def ingest_document(file_path: str, filename: str):
    """
    Ingests a document, splits it into chunks, and stores it in the vector DB.
    """
    ext = filename.split(".")[-1].lower()
    
    if ext == "pdf":
        loader = PyMuPDFLoader(file_path)
        documents = loader.load()
    elif ext == "csv":
        loader = CSVLoader(file_path)
        documents = loader.load()
    elif ext in ["txt", "md"]:
        # Try utf-8 first, fallback to latin-1
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
        except Exception:
            loader = TextLoader(file_path, encoding='latin-1')
            documents = loader.load()
    elif ext == "docx":
        documents = load_docx(file_path)
    elif ext == "json":
        documents = load_json(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
        
    # Chunking strategy — larger chunks for better context
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=300,
        length_function=len
    )
    
    chunks = text_splitter.split_documents(documents)
    
    # Add metadata
    for chunk in chunks:
        chunk.metadata["source_filename"] = filename
        
    # Add to vector store in batches to avoid memory issues
    batch_size = 50
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        vector_store.add_documents(batch)
    
    return len(chunks)


def delete_vectors_by_filename(filename: str):
    """Delete all vector chunks associated with a specific filename from ChromaDB."""
    try:
        collection = vector_store._collection
        # Get all IDs where source_filename matches
        results = collection.get(
            where={"source_filename": filename}
        )
        if results and results["ids"]:
            collection.delete(ids=results["ids"])
            return len(results["ids"])
        return 0
    except Exception as e:
        print(f"Error deleting vectors for {filename}: {e}")
        return 0
